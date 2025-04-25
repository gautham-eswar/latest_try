#!/usr/bin/env python3
"""
Resume Processing System Test Script
This script validates the resume processing implementation with test data and specific component tests.
"""

import os
import sys
import time
import unittest
import json
import io
import tempfile
import random
import string
import logging
from datetime import datetime
from pathlib import Path
from contextlib import contextmanager
from unittest.mock import patch, MagicMock

# Set up basic logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('test_resume_processing')

# Constants for test configuration
TEST_DIR = Path(tempfile.mkdtemp(prefix='resume_test_'))
SUCCESS_THRESHOLD = 70  # Success rate below this percentage will fail tests


class TimedTestResult(unittest.TextTestResult):
    """Custom test result class that tracks execution time for each test."""
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.test_timings = {}
        self._start_time = 0
        
    def startTest(self, test):
        self._start_time = time.time()
        super().startTest(test)
        
    def addSuccess(self, test):
        elapsed = time.time() - self._start_time
        self.test_timings[test.id()] = elapsed
        logger.info(f"✅ PASS: {test.id()} ({elapsed:.3f}s)")
        super().addSuccess(test)
        
    def addFailure(self, test, err):
        elapsed = time.time() - self._start_time
        self.test_timings[test.id()] = elapsed
        logger.error(f"❌ FAIL: {test.id()} ({elapsed:.3f}s)")
        logger.error(f"Error: {err[1]}")
        super().addFailure(test, err)
        
    def addError(self, test, err):
        elapsed = time.time() - self._start_time
        self.test_timings[test.id()] = elapsed
        logger.error(f"⚠️ ERROR: {test.id()} ({elapsed:.3f}s)")
        logger.error(f"Error: {err[1]}")
        super().addError(test, err)


class TimedTextTestRunner(unittest.TextTestRunner):
    """Custom test runner that uses our TimedTestResult."""
    
    def __init__(self, *args, **kwargs):
        kwargs['resultclass'] = TimedTestResult
        super().__init__(*args, **kwargs)
        
    def run(self, test):
        result = super().run(test)
        total_time = sum(result.test_timings.values())
        print(f"\nTotal test execution time: {total_time:.3f} seconds")
        return result


@contextmanager
def measure_time(operation_name):
    """Context manager to measure and log execution time of operations."""
    start_time = time.time()
    try:
        yield
    finally:
        elapsed = time.time() - start_time
        logger.info(f"{operation_name} completed in {elapsed:.3f} seconds")


def generate_sample_pdf():
    """Generate a minimal sample PDF resume for testing."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = TEST_DIR / "sample_resume.pdf"
        
        # Create a PDF with minimal resume content
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        
        # Add resume content
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, 750, "Sample Test Resume")
        
        c.setFont("Helvetica-Bold", 12)
        c.drawString(72, 720, "Contact Information")
        c.setFont("Helvetica", 10)
        c.drawString(72, 705, "Email: test@example.com")
        c.drawString(72, 690, "Phone: 555-123-4567")
        
        c.setFont("Helvetica-Bold", 12)
        c.drawString(72, 660, "Skills")
        c.setFont("Helvetica", 10)
        c.drawString(72, 645, "Python, JavaScript, Flask, Data Analysis, Machine Learning")
        
        c.setFont("Helvetica-Bold", 12)
        c.drawString(72, 615, "Experience")
        c.setFont("Helvetica-Bold", 10)
        c.drawString(72, 600, "Software Engineer - ABC Company")
        c.setFont("Helvetica", 10)
        c.drawString(72, 585, "2020-2023")
        c.drawString(72, 570, "Developed web applications using Python and Flask")
        c.drawString(72, 555, "Implemented machine learning algorithms for data analysis")
        
        c.setFont("Helvetica-Bold", 12)
        c.drawString(72, 525, "Education")
        c.setFont("Helvetica-Bold", 10)
        c.drawString(72, 510, "Bachelor of Science in Computer Science")
        c.setFont("Helvetica", 10)
        c.drawString(72, 495, "XYZ University, 2016-2020")
        
        c.save()
        logger.info(f"Generated sample PDF resume at {pdf_path}")
        return pdf_path
    except ImportError:
        logger.error("Could not import reportlab - cannot generate test PDF")
        return None


def generate_sample_docx():
    """Generate a minimal sample DOCX resume for testing."""
    try:
        from docx import Document
        from docx.shared import Pt
        
        docx_path = TEST_DIR / "sample_resume.docx"
        doc = Document()
        
        # Add title
        title = doc.add_heading("Sample Test Resume", 0)
        
        # Add contact information
        doc.add_heading("Contact Information", level=2)
        contact = doc.add_paragraph()
        contact.add_run("Email: test@example.com\nPhone: 555-123-4567")
        
        # Add skills
        doc.add_heading("Skills", level=2)
        skills = doc.add_paragraph()
        skills.add_run("Python, JavaScript, Flask, Data Analysis, Machine Learning")
        
        # Add experience
        doc.add_heading("Experience", level=2)
        company = doc.add_paragraph()
        company.add_run("Software Engineer - ABC Company\n").bold = True
        company.add_run("2020-2023\n")
        company.add_run("• Developed web applications using Python and Flask\n")
        company.add_run("• Implemented machine learning algorithms for data analysis")
        
        # Add education
        doc.add_heading("Education", level=2)
        education = doc.add_paragraph()
        education.add_run("Bachelor of Science in Computer Science\n").bold = True
        education.add_run("XYZ University, 2016-2020")
        
        # Save the document
        doc.save(str(docx_path))
        logger.info(f"Generated sample DOCX resume at {docx_path}")
        return docx_path
    except ImportError:
        logger.error("Could not import python-docx - cannot generate test DOCX")
        return None


def generate_job_description():
    """Generate a sample job description file with relevant keywords."""
    job_desc_path = TEST_DIR / "sample_job_description.txt"
    job_description = """
    Job Title: Senior Python Developer
    
    Company: Tech Innovations Inc.
    
    Job Description:
    We are seeking an experienced Python developer with strong Flask background to join our team.
    The ideal candidate will have experience developing web applications, implementing data analysis
    algorithms, and working with cloud services.
    
    Requirements:
    - 3+ years of experience with Python development
    - Strong knowledge of Flask, FastAPI, or Django
    - Experience with RESTful API design and implementation
    - Familiarity with machine learning libraries such as TensorFlow or PyTorch
    - Database knowledge (PostgreSQL, MongoDB)
    - Experience with AWS or other cloud providers
    - BS in Computer Science or related field
    - Excellent problem-solving and communication skills
    
    Responsibilities:
    - Design and implement new features for our data processing platform
    - Optimize existing algorithms for better performance
    - Work with the ML team to integrate machine learning models
    - Maintain and improve the existing codebase
    - Participate in code reviews and architectural discussions
    
    Benefits:
    - Competitive salary and benefits package
    - Flexible work hours
    - Remote work options
    - Professional development opportunities
    """
    
    with open(job_desc_path, "w") as f:
        f.write(job_description)
        
    logger.info(f"Generated sample job description at {job_desc_path}")
    return job_desc_path


def setup_test_environment():
    """Set up the test environment with sample files."""
    logger.info("Setting up test environment...")
    
    environment = {
        "pdf_resume": generate_sample_pdf(),
        "docx_resume": generate_sample_docx(),
        "job_description": generate_job_description(),
        "test_dir": TEST_DIR
    }
    
    # Create output directory for generated files
    output_dir = TEST_DIR / "output"
    output_dir.mkdir(exist_ok=True)
    environment["output_dir"] = output_dir
    
    # Create a malformed PDF for error testing
    malformed_pdf = TEST_DIR / "malformed.pdf"
    with open(malformed_pdf, "w") as f:
        f.write("This is not a real PDF file, just text with a PDF extension.")
    environment["malformed_pdf"] = malformed_pdf
    
    logger.info(f"Test environment set up at {TEST_DIR}")
    return environment


def cleanup_test_environment(env):
    """Clean up the test environment."""
    import shutil
    try:
        shutil.rmtree(env["test_dir"])
        logger.info(f"Cleaned up test directory {env['test_dir']}")
    except Exception as e:
        logger.error(f"Failed to clean up test directory: {e}")


class ResumeProcessingTestCase(unittest.TestCase):
    """Test case for the resume processing system."""
    
    @classmethod
    def setUpClass(cls):
        """Set up the test environment once for all tests."""
        cls.env = setup_test_environment()
        
        # Import necessary modules for testing
        try:
            # Add directory to path if needed
            sys.path.insert(0, os.path.abspath("./resume-o"))
            
            # Import application and modules
            from app import create_app
            from database import Database
            # Import other required modules
            logger.info("Successfully imported required modules")
            
            # Create a test Flask app
            cls.app = create_app(testing=True)
            cls.client = cls.app.test_client()
            cls.app_context = cls.app.app_context()
            cls.app_context.push()
            
        except ImportError as e:
            logger.error(f"Failed to import required modules: {e}")
            raise
    
    @classmethod
    def tearDownClass(cls):
        """Clean up after all tests."""
        if hasattr(cls, 'app_context'):
            cls.app_context.pop()
        cleanup_test_environment(cls.env)
    
    def test_010_app_initialization(self):
        """Test that the application initializes correctly."""
        with measure_time("App initialization test"):
            self.assertIsNotNone(self.app, "Flask app should be created")
            self.assertTrue(self.app.config['TESTING'], "App should be in testing mode")
    
    def test_020_file_upload_endpoint(self):
        """Test the file upload endpoint."""
        with measure_time("File upload test"):
            if not self.env.get("pdf_resume"):
                self.skipTest("PDF test file not available")
                
            with open(self.env["pdf_resume"], "rb") as pdf:
                response = self.client.post(
                    "/api/upload",
                    data={
                        "file": (pdf, "sample_resume.pdf"),
                    },
                    content_type="multipart/form-data"
                )
            
            self.assertEqual(response.status_code, 200, f"Upload failed with status {response.status_code}")
            data = json.loads(response.data)
            self.assertIn("resume_id", data, "Response should contain resume_id")
            self.assertIn("parsed_data", data, "Response should contain parsed_data")
            
            # Store resume_id for later tests
            self.__class__.resume_id = data["resume_id"]
    
    def test_030_resume_parsing(self):
        """Test resume parsing functionality."""
        with measure_time("Resume parsing test"):
            if not self.env.get("pdf_resume"):
                self.skipTest("PDF test file not available")
            
            try:
                # Import the parser directly
                from resume_parser import parse_resume
                
                result = parse_resume(self.env["pdf_resume"])
                self.assertIsNotNone(result, "Parsed result should not be None")
                self.assertIn("contact", result, "Parsed data should contain contact information")
                self.assertIn("skills", result, "Parsed data should contain skills")
                self.assertIn("experience", result, "Parsed data should contain experience")
                self.assertIn("education", result, "Parsed data should contain education")
            except ImportError:
                # Try through API if direct import fails
                with open(self.env["pdf_resume"], "rb") as pdf:
                    response = self.client.post(
                        "/api/upload",
                        data={
                            "file": (pdf, "sample_resume.pdf"),
                        },
                        content_type="multipart/form-data"
                    )
                
                self.assertEqual(response.status_code, 200, f"Upload failed with status {response.status_code}")
                data = json.loads(response.data)
                self.assertIn("parsed_data", data, "Response should contain parsed_data")
    
    def test_040_keyword_extraction(self):
        """Test keyword extraction from job description."""
        with measure_time("Keyword extraction test"):
            if not self.env.get("job_description"):
                self.skipTest("Job description test file not available")
            
            try:
                # Import the keyword extractor directly
                from keyword_extractor import extract_keywords
                
                with open(self.env["job_description"], "r") as f:
                    job_text = f.read()
                
                keywords = extract_keywords(job_text)
                self.assertIsNotNone(keywords, "Keywords should not be None")
                self.assertGreater(len(keywords), 0, "Should extract at least some keywords")
                
                # Check for expected keywords in the sample job description
                expected_keywords = ["python", "flask", "machine learning", "data"]
                for keyword in expected_keywords:
                    self.assertTrue(
                        any(keyword.lower() in kw.lower() for kw in keywords),
                        f"Expected keyword '{keyword}' not found in extracted keywords"
                    )
            except ImportError:
                # Try through API if direct import fails
                with open(self.env["job_description"], "r") as f:
                    job_text = f.read()
                
                response = self.client.post(
                    "/api/extract-keywords",
                    json={"text": job_text}
                )
                
                self.assertEqual(response.status_code, 200, f"Keyword extraction failed with status {response.status_code}")
                data = json.loads(response.data)
                self.assertIn("keywords", data, "Response should contain keywords")
                self.assertGreater(len(data["keywords"]), 0, "Should extract at least some keywords")
    
    def test_050_semantic_matching(self):
        """Test semantic matching between resume and job description."""
        with measure_time("Semantic matching test"):
            if not hasattr(self.__class__, "resume_id"):
                self.skipTest("Resume ID from previous test not available")
            
            if not self.env.get("job_description"):
                self.skipTest("Job description test file not available")
            
            with open(self.env["job_description"], "r") as f:
                job_text = f.read()
            
            response = self.client.post(
                "/api/match",
                json={
                    "resume_id": self.__class__.resume_id,
                    "job_description": job_text
                }
            )
            
            self.assertEqual(response.status_code, 200, f"Matching failed with status {response.status_code}")
            data = json.loads(response.data)
            self.assertIn("match_score", data, "Response should contain match_score")
            self.assertIn("missing_keywords", data, "Response should contain missing_keywords")
            self.assertIn("matching_keywords", data, "Response should contain matching_keywords")
    
    def test_060_resume_enhancement(self):
        """Test resume enhancement based on job description."""
        with measure_time("Resume enhancement test"):
            if not hasattr(self.__class__, "resume_id"):
                self.skipTest("Resume ID from previous test not available")
            
            if not self.env.get("job_description"):
                self.skipTest("Job description test file not available")
            
            with open(self.env["job_description"], "r") as f:
                job_text = f.read()
            
            response = self.client.post(
                "/api/optimize",
                json={
                    "resume_id": self.__class__.resume_id,
                    "job_description": job_text,
                    "enhancement_level": "moderate"
                }
            )
            
            self.assertEqual(response.status_code, 200, f"Enhancement failed with status {response.status_code}")
            data = json.loads(response.data)
            self.assertIn("enhanced_resume_id", data, "Response should contain enhanced_resume_id")
            self.assertIn("enhancements", data, "Response should contain enhancements")
            
            # Store enhanced_resume_id for later tests
            self.__class__.enhanced_resume_id = data["enhanced_resume_id"]
    
    def test_070_pdf_generation(self):
        """Test PDF generation from enhanced resume."""
        with measure_time("PDF generation test"):
            if not hasattr(self.__class__, "enhanced_resume_id"):
                self.skipTest("Enhanced resume ID from previous test not available")
            
            response = self.client.get(
                f"/api/download/{self.__class__.enhanced_resume_id}/pdf"
            )
            
            self.assertEqual(response.status_code, 200, f"PDF download failed with status {response.status_code}")
            self.assertEqual(response.mimetype, "application/pdf", "Response should be a PDF")
            
            # Save the generated PDF for inspection
            output_path = self.env["output_dir"] / "enhanced_resume.pdf"
            with open(output_path, "wb") as f:
                f.write(response.data)
            
            logger.info(f"Generated enhanced PDF saved to {output_path}")
    
    def test_080_docx_generation(self):
        """Test DOCX generation from enhanced resume."""
        with measure_time("DOCX generation test"):
            if not hasattr(self.__class__, "enhanced_resume_id"):
                self.skipTest("Enhanced resume ID from previous test not available")
            
            response = self.client.get(
                f"/api/download/{self.__class__.enhanced_resume_id}/docx"
            )
            
            self.assertEqual(response.status_code, 200, f"DOCX download failed with status {response.status_code}")
            self.assertEqual(
                response.mimetype, 
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                "Response should be a DOCX"
            )
            
            # Save the generated DOCX for inspection
            output_path = self.env["output_dir"] / "enhanced_resume.docx"
            with open(output_path, "wb") as f:
                f.write(response.data)
            
            logger.info(f"Generated enhanced DOCX saved to {output_path}")
    
    def test_090_error_handling_malformed_file(self):
        """Test error handling for malformed files."""
        with measure_time("Malformed file error handling test"):
            if not self.env.get("malformed_pdf"):
                self.skipTest("Malformed PDF test file not available")
                
            with open(self.env["malformed_pdf"], "rb") as pdf:
                response = self.client.post(
                    "/api/upload",
                    data={
                        "file": (pdf, "malformed.pdf"),
                    },
                    content_type="multipart/form-data"
                )
            
            self.assertNotEqual(response.status_code, 200, "Should reject malformed file")
            self.assertEqual(response.status_code, 400, "Should return 400 Bad Request for malformed file")
    
    def test_100_error_handling_missing_params(self):
        """Test error handling for missing parameters."""
        with measure_time("Missing parameters error handling test"):
            response = self.client.post(
                "/api/optimize",
                json={
                    # Missing resume_id
                    "job_description": "Sample job description"
                }
            )
            
            self.assertNotEqual(response.status_code, 200, "Should reject request with missing parameters")
            self.assertEqual(response.status_code, 400, "Should return 400 Bad Request for missing parameters")
    
    def test_110_pipeline_end_to_end(self):
        """Test the complete resume processing pipeline end to end."""
        with measure_time("End-to-end pipeline test"):
            if not self.env.get("pdf_resume") or not self.env.get("job_description"):
                self.skipTest("Test files not available")
            
            # Step 1: Upload resume
            with open(self.env["pdf_resume"], "rb") as pdf:
                response1 = self.client.post(
                    "/api/upload",
                    data={
                        "file": (pdf, "sample_resume.pdf"),
                    },
                    content_type="multipart/form-data"
                )
            
            self.assertEqual(response1.status_code, 200, f"Upload failed with status {response1.status_code}")
            data1 = json.loads(response1.data)
            resume_id = data1["resume_id"]
            
            # Step 2: Get job description
            with open(self.env["job_description"], "r") as f:
                job_text = f.read()
            
            # Step 3: Optimize resume
            response2 = self.client.post(
                "/api/optimize",
                json={
                    "resume_id": resume_id,
                    "job_description": job_text,
                    "enhancement_level": "moderate"
                }
            )
            
            self.assertEqual(response2.status_code, 200, f"Enhancement failed with status {response2.status_code}")
            data2 = json.loads(response2.data)
            enhanced_resume_id = data2["enhanced_resume_id"]
            
            # Step 4: Download in PDF format
            response3 = self.client.get(
                f"/api/download/{enhanced_resume_id}/pdf"
            )
            
            self.assertEqual(response3.status_code, 200, f"PDF download failed with status {response3.status_code}")
            
            # Step 5: Download in DOCX format
            response4 = self.client.get(
                f"/api/download/{enhanced_resume_id}/docx"
            )
            
            self.assertEqual(response4.status_code, 200, f"DOCX download failed with status {response4.status_code}")
            
            # Step 6: Download in LaTeX format
            response5 = self.client.get(
                f"/api/download/{enhanced_resume_id}/latex"
            )
            
            self.assertEqual(response5.status_code, 200, f"LaTeX download failed with status {response5.status_code}")
            
            logger.info("Successfully completed end-to-end pipeline test")


if __name__ == "__main__":
    try:
        print("=" * 80)
        print(f"Resume Processing System Test - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("=" * 80)
        
        # Check if test environment can be set up
        test_env = setup_test_environment()
        if not test_env["pdf_resume"] and not test_env["docx_resume"]:
            logger.error("Failed to generate test files. Please install reportlab and python-docx packages.")
            sys.exit(1)
        cleanup_test_environment(test_env)
        
        # Run the tests
        test_suite = unittest.TestLoader().loadTestsFromTestCase(ResumeProcessingTestCase)
        result = TimedTextTestRunner(verbosity=2).run(test_suite)
        
        # Display summary
        print("\n" + "=" * 80)
        print(f"SUMMARY: Ran {result.testsRun} tests")
        print(f"SUCCESS: {result.testsRun - len(result.failures) - len(result.errors)} tests")
        print(f"FAILURES: {len(result.failures)} tests")
        print(f"ERRORS: {len(result.errors)} tests")
        
        # Exit with appropriate code
        sys.exit(len(result.failures) + len(result.errors))
        
    except KeyboardInterrupt:
        print("\nTest interrupted by user")
        sys.exit(1)
    except Exception as e:
        print(f"\nUnexpected error during test execution: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1) 